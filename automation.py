import os
from datetime import datetime
from docx import Document

folder_path = r'C:\Users\Hp\Desktop\marium uni work\CodeAlphaDemo'

if not os.path.exists(folder_path):
    os.makedirs(folder_path)

doc = Document()

doc.add_heading('Data Communication and Networks lab task', 0)
doc.add_paragraph('Name: Marium Haris')
doc.add_paragraph('Student ID: 15084')


file_name = f"dcn lab task.docx"
file_path = os.path.join(folder_path, file_name)

doc.save(file_path)

print(f"Word document '{file_name}' created and saved in {folder_path}.")
